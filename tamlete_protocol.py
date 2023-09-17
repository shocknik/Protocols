import re
import os
import docx
import datetime
from setting import standarts, list_head_SI_IO, list_mean_SI_IO, list_head_test_table
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Cm, Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION
from borders import set_cell_border, add_page_number
from backend_read import get_cable_mark,\
    get_specifications,\
    get_name_specifications,\
    get_list_text_par,\
    create_marker_list,\
    change_font, border_form,\
    func_union_cells,\
    filling_table_heads,\
    filling_table_heads_all

doc = Document()
num_page = 21

"""Формирование альбомной формы, шрифта и отступов """
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

title = doc.sections[0]
title.orientation = WD_ORIENTATION.LANDSCAPE
title.page_height = Cm(21.0)
title.page_width = Cm(29.7)

# левое поле в миллиметрах
title.left_margin = Mm(15)
title.right_margin = Mm(13)
title.top_margin = Mm(15)
title.bottom_margin = Mm(13)


title_table = doc.add_table(rows=7, cols=3)

border_form(7, 3, title_table)        

cell_header = title_table.cell(0, 1)
cell_header.paragraphs[0].add_run('Общество с ограниченной ответственностью НИЦ «Кабель-Тест»\n(ООО НИЦ «Кабель-Тест»)').font.size = Pt(14)
cell_header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_header.width = Cm(23)
cell_header.paragraphs[0].paragraph_format.space_after = Pt(0)


title_table.cell(1, 2).merge(title_table.cell(2, 2))
cell_seo = title_table.cell(1, 2)
cell_seo.paragraphs[0].add_run("УТВЕРЖДАЮ\n").bold = True
cell_seo.paragraphs[0].add_run('Генеральный директор\n\
ООО НИЦ «Кабель-Тест»-\n\
Руководитель ИЦ\n\
___________ В.И. Видяев\n\
«___» августа 2023 г.'
).font.size = Pt(12)
cell_seo.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


title_table.cell(2, 0).merge(title_table.cell(2, 1))
cell_adress = title_table.cell(2, 0)
cell_adress.paragraphs[0].add_run('Юридический адрес ООО НИЦ «Кабель-Тест»:\n\
123290, г. Москва, Магистральный 1-й туп., д. 5А, комн. 132Л\n\
Адрес места нахождения и осуществления деятельности:\n\
107497, г. Москва, ул. Бирюсинка, д. 6, корп. 1-5, 6, 7, 9А\n\
Телефон: +7 495 603-06-55, e-mail: info@cable-test.ru\n\
Уникальный номер в реестре аккредитованных лиц: РОСС RU.0001.21КБ32').font.size = Pt(12)
cell_adress.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cell_adress.paragraphs[0].paragraph_format.line_spacing = Pt(12)
cell_adress.width = Cm(15)


cell_header_2 = title_table.cell(1, 1)
cell_header_2.paragraphs[0].add_run('ИСПЫТАТЕЛЬНЫЙ ЦЕНТР').bold = True
cell_header_2.paragraphs[0].paragraph_format.space_before = Pt(0)
cell_header_2.paragraphs[0].paragraph_format.space_after = Cm(1)
cell_header_2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
set_cell_border(cell_header_2, top={"sz": 12, "val": "single", "color": "black", "space": "0"})





title_table.cell(3, 0).merge(title_table.cell(3, 1))
title_table.cell(3, 1).merge(title_table.cell(3, 2))
cell_prot = title_table.cell(3, 1)
cell_prot.paragraphs[0].add_run("ПРОТОКОЛ №            \n").bold = True
cell_prot.paragraphs[0].add_run("от              \n\
приемочных испытаний")
cell_prot.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cell_prot.paragraphs[0].paragraph_format.space_after = Pt(0)


cell_discription = title_table.cell(4, 1)
cell_discription.paragraphs[0].add_run("кабеля                                      марки ")
cell_discription.paragraphs[0].add_run("{},\n".format(get_cable_mark()[0])).bold = True
cell_discription.paragraphs[0].add_run("изготовленного ООО НПП «СПЕЦКАБЕЛЬ»\
    на соответствие требованиям {} {}".format(get_specifications(), get_name_specifications()))
cell_discription.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_HI
cell_discription.paragraphs[0].paragraph_format.line_spacing = Pt(12)
cell_discription.paragraphs[0].paragraph_format.space_after = Pt(24)



title_table.cell(5, 0).merge(title_table.cell(5, 1))
title_table.cell(5, 1).merge(title_table.cell(5, 2))
cell_remarks = title_table.cell(5, 0)
cell_remarks.paragraphs[0].add_run("1 Листов всего {}\n".format(num_page))
cell_remarks.paragraphs[0].add_run("2 Результаты испытаний распространяются только на предоставленный (е) заказчиком образец (цы).\n\
3 Протокол испытаний не может быть частично или полностью воспроизведен без письменного разрешения Испытательного центра.")
cell_remarks.paragraphs[0].paragraph_format.space_after = Cm(2)    

set_cell_border(title_table.cell(5, 0), end={"sz": 12, "val": "single", "color": "black", "space": "0"})
set_cell_border(title_table.cell(3, 0), end={"sz": 12, "val": "single", "color": "black", "space": "0"})
set_cell_border(title_table.cell(6, 0), end={"sz": 12, "val": "single", "color": "black", "space": "0"})

title_table.cell(6, 0).merge(title_table.cell(6, 1))
title_table.cell(6, 1).merge(title_table.cell(6, 2))
cell_moscow = title_table.cell(6, 0)
cell_moscow.paragraphs[0].add_run("Москва\n2023")
cell_moscow.paragraphs[0].paragraph_format.space_after = Pt(0)
cell_moscow.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


doc.add_section()

data_protocol = doc.sections[1]

doc.add_paragraph()

head_1 = doc.add_paragraph().add_run('1 Основание для проведения испытаний')
head_1.font.size = Pt(12)
head_1.bold = True
head_1_1 = doc.add_paragraph().add_run('Программа типовых испытаний кабелей марок\
{} {}, изготовленных и представленных на испытания ООО НПП «Спецкабель»\
на соответствие требованиям {} {}'.format(
    get_cable_mark()[0],
    get_cable_mark()[1],
    get_specifications(),
    get_name_specifications()
))

head_2 = doc.add_paragraph().add_run('2 Информация о заказчике и изготовителе')
head_2.font.size = Pt(12)
head_2.bold = True
head_2_1 = doc.add_paragraph().add_run('ООО НПП «Спецкабель» (ИНН 7701165130)\n\
•	юридический адрес и фактический адрес: 107497, г. Москва, ул. Бирюсинка, д. 6, к. 1-5, пом. XVI, ком. 15\n\
•	телефон.: +7 (495) 134-2-134n\n\
•	е-mail: info@spetskabel.ru')

head_3 = doc.add_paragraph().add_run('3 Наименование объекта испытаний')
head_3.font.size = Pt(12)
head_3.bold = True

head_3_1 = doc.add_paragraph().add_run('Отбор образцов проведен ООО НПП «Спецкабель»\n\
Количество образцов 1. Длина образцов 100 м\n\
Номер партии:            \n\
Образцы представленны на испытания заказчиком. Состояние образцов - без видимых повреждений.\n\
На испытания представлен образцы кабеля       \n\
Образцы изготовленны по {} {}'.format(get_specifications(), get_name_specifications()))
head_3_1 = doc.add_paragraph().add_run('Предоставленный(-е) на испытания образец(-цы) идентифицирован(-ы)\
как объект испытаний путем сравнения основных характеристик изделия,\
указанных в сопроводительной и технической документации с фактическими данными\
на образце(-ах).\n\
Объекту испытаний присвоен уникальный идентификационный номер ID 0000')

head_4 = doc.add_paragraph().add_run('4	Даты проведения испытаний ')
head_4.font.size = Pt(12)
head_4.bold = True
head_4_1 = doc.add_paragraph().add_run('Начало испытаний:\nОкончание испытаний:')

head_5 = doc.add_paragraph().add_run('5 Цель испытаний')
head_5.font.size = Pt(12)
head_5.bold = True
head_5_1 = doc.add_paragraph().add_run('Определение соответствия образцов кабеля\
требованиям {} {}'.format(get_specifications(), get_name_specifications()))

head_6 = doc.add_paragraph().add_run('6 Условия окружающей среды при проведении испытаний')
head_6.font.size = Pt(12)
head_6.bold = True
head_6_1 = doc.add_paragraph().add_run('Испытания проводились при внешних условиях окружающей среды:\n\
- температура -  °С;\n\
- относительная влажность воздуха -  %;\n\
- атмосферное давление -  кПа.')

head_7 = doc.add_paragraph().add_run('7 Методы испытаний')
head_7.font.size = Pt(12)
head_7.bold = True
head_7_1 = doc.add_paragraph().add_run('Методы испытаний в соответствии с требованиями:\n')

"""Создаем список стандартов"""
for standart in standarts:
    st = doc.add_paragraph(standart, 'List Paragraph')
    unordered = "1" # задаем номер и типа маркера/отступа
    create_marker_list(st, unordered)


head_8 = doc.add_paragraph().add_run('8 Испытательное оборудование и средства измерений')
head_8.font.size = Pt(12)
head_8.bold = True
head_8_1 = doc.add_paragraph().add_run('9 Применяемые испытательное оборудование (ИО) и средства измерений (СИ) представлены в таблице 1.')
head_8_2 = doc.add_paragraph('Таблица 1')
head_8_2.paragraph_format.space_after = Pt(0)
change_font(head_8_2.runs[0])


"""Создаем таблицу с СИ и ИО"""
table_SI_IO = doc.add_table(2, 7)
table_SI_IO.style = 'Table Grid'
filling_table_heads_all(table_SI_IO, list_head_SI_IO)

"""Заполняем таблицу приборами"""
for row in list_mean_SI_IO:
    cells = table_SI_IO.add_row().cells
    for i, list_mean_SI_IO in enumerate(row):
        cells[i].text = str(list_mean_SI_IO)
    
head_9 = doc.add_paragraph('10 Результаты измерений\n')
head_9.runs[0].font.size = Pt(12)
head_9.runs[0].bold = True
head_9.add_run('10.1 Результаты испытаний на соответствие техническим требованиям представлены в таблице 2.')
head_9.runs[1].font.size = Pt(12)


"""Формирование таблицы испытаний"""
test_table = doc.add_table(rows=2, cols=7)
cells_union = {
    "1": [0, 1],
    "2": [0, 2],
    "3": [0, 3],
    "4": [0, 4],
    "5": [0, 0],
    "6": [1, 0],
    "7": [0, 5],
    "8": [1, 5],
    "9": [0, 6],
    "10": [1, 6],
}

border_form(2, 7, test_table, border="double", sz=6)
func_union_cells(test_table, **cells_union)
test_table.style = 'Table Grid'
filling_table_heads_all(test_table, list_head_test_table)







doc.save('title.docx')