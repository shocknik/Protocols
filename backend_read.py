"""
TODO
0 Попытаться научиться распаковывать ДОК-файлы и парсить xml
import xml.etree.ElementTree as ET
import lxml
import zipfile
from bs4 import BeautifulSoup

doc_zip = zipfile.ZipFile("Программы квалификационных испытаний\\Программы квалификационных испытаний\\020\\Программа типовых испытаний ТУ 16.К99-020-2009.docx")
doc_xml = doc_zip.read('word/document.xml')

soup_xml = BeautifulSoup(doc_xml, features='xml')
pretty_xml = soup_xml.prettify()
root = ET.fromstring(pretty_xml)
print(root)
1 Создать шаблон протокола испытаний
    - автоматизировать генерацию шаблона с присвоением номер +1
    - делать запрос на чтение программы
    - создавать таблицу приборов
    - создавать таблицу испытаний
2 Изучить программы испытаний
    - выявить ключевые слова, к которым могут относиться:
        - марка кабеля
        - номер ТУ
        - наименование ТУ
    - научится переходить к таблице и читать ее
    - читать испытания в таблице и собирать их в словарь или лист
    - заполнять по полученому словарю таблицу в протоколе
3 Структура проекта:
    3.1 main.py (главный исполняемый файл, который генерит запросы)
    3.2 backend.py (файл содержащий функции)
    3.3 model.py (файл который реализует логику, используя функции backend.py)
    3.4 view.py (отображает информацию через терминал)
    3.5 controller.py (осуществляет взаимодействие между model.py и veiw.py)
"""
import json
import re
import os
import docx
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.text import font
from docx.oxml.ns import qn
from borders import set_cell_border
from setting import *

document = Document()

def get_paths_filesdoc() -> list:
    """Получить пути всех файлов doc в директории"""
    paths = []
    folder = os.getcwd()
    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.endswith('docx') and not file.startswith('~'):
                paths.append(os.path.join(root, file))
    return paths


doc = Document('Программы квалификационных испытаний\\Программы квалификационных испытаний\\064\\Программа типовых испытаний ФЖТК.357400.064ТУ.docx')

list_par: list = doc.paragraphs # список объектов параграфов


def get_list_text_par() -> list:
    """получаем список тесктов параграфов"""
    list_text_par: list = []
    for par in list_par:
        list_text_par.append(par.text)
    return list_text_par

        
def get_cable_mark() -> list:
    """Получение списка марок кабелей из программы"""
    for par in get_list_text_par():
        result = re.findall(r'[0-9А-Яа-я\S]+\b\s\dх\dх\d[,]*\d*', par) # марка кабеля
        if len(result) > 0:
            return result # Печать списка марок
       
        
def get_specifications() -> str:
    """Получение ТУ на кабели из программы"""
    for par in get_list_text_par():
        result = re.findall(r'[ФЖТК]+\W[0-9]+\W[0-9ТУ]+|[ТУ]+[\s]\d{1,3}\S[К1-9\S]+', par)
        if 0 < len(result) < 2:
            return result[0]


def get_name_specifications() -> str:
    """Получение названия ТУ"""
    for par in get_list_text_par():
        result = re.findall(r'\«Кабели\s*[а-яА-Я\s,0-9]+\.\s[Техническиеусловия\s]+\»', par) # название ТУ
        if 0 < len(result) < 2:
            return result[0]


def get_list_par_from_tables() -> list:
    """Получает список параметров из таблицы в программе"""
    list_par = []
    for table in doc.tables:
        for i in range(2, len(table.rows)):
            par = ' '.join(table.cell(i, 1).text.split())
            list_par.append(par)
    return list_par
            
            
def get_list_requarements() -> list:
    """Полчает список пунктов требований"""
    list_req = []
    for table in doc.tables:
        for i in range(2, len(table.rows)):
            list_req.append(table.cell(i, 2).text)
    return list_req
                

def get_list_methods() -> list:
    """Полчает список пунктов методов"""
    list_methods = []
    for table in doc.tables:
        for i in range(2, len(table.rows)):
            list_methods.append(table.cell(i, 3).text)
    return list_methods
            
def create_marker_list(paragraph, list_type):
    "функция создания маркированного списка"
    p = paragraph._p # доступ к элементам xml параграфов
    pPr = p.get_or_add_pPr() # доступ к свойствам параграфов
    numPr = OxmlElement('w:numPr') # создать элемент свойств числа
    numId =OxmlElement('w:numId') # создать элемент маркера и тип
    numId.set(qn('w:val'), list_type) # тип списка настройка отступов
    numPr.append(numId) # добавить тип маркера в список свойств числа
    pPr.append(numPr) # добавить свойства числа в абзац
    
    
def change_font(obj_run):
    """Функция для формирования межсимвольного интервала"""
    obj_r = obj_run._r
    rPr = obj_r.get_or_add_rPr()
    space = OxmlElement('w:spacing')
    space.set(qn('w:val'), '20')
    rPr.append(space)
    
def border_form(rows: int, cols: int, table, border="single", sz=12):
    """Функция, которая рисует рамку для указанной таблице, вокруг указанных яйчеек"""
    for row in range(0, rows):
        for col in range(0, cols):
            if row == 0:
                set_cell_border(table.cell(row, col), top={"sz": sz, "val": border, "color": "black", "space": "0"})
            if row == rows - 1: 
                set_cell_border(table.cell(row, col), bottom={"sz": sz, "val": border, "color": "black", "space": "0"})
            if col == 0:
                set_cell_border(table.cell(row, col), start={"sz": sz, "val": border, "color": "black", "space": "0"})
            if col == cols - 1:
                set_cell_border(table.cell(row, col), end={"sz": sz, "val": border, "color": "black", "space": "0"})

def border_around_cell(cell, border="double", sz=6):
    set_cell_border(cell, top={"sz": sz, "val": border, "color": "black", "space": "0"})
    set_cell_border(cell, bottom={"sz": sz, "val": border, "color": "black", "space": "0"})
    set_cell_border(cell, start={"sz": sz, "val": border, "color": "black", "space": "0"})
    set_cell_border(cell, end={"sz": sz, "val": border, "color": "black", "space": "0"})
    """Функция рисующая границы вокруг ячейки"""


def table_inner_border_vertical(rows: int, cols: int, table, border="single", sz=12, vert=True):
    """
    Функция которая строит внутренние вертикальные границы.
    Если vert = False, то рисует еще и горизонтальные границы
    
    """
    for row in range(0, rows):
        for col in range(0, cols):
            set_cell_border(table.cell(row, col), start={"sz": sz, "val": border, "color": "black", "space": "0"})
            if vert is False:
                set_cell_border(table.cell(row, col), bottom={"sz": sz, "val": border, "color": "black", "space": "0"})
           

def func_union_cells(table, **cells):
    """
    Функция, которая объединяет яйчейки в таблице.
    На вход принимает имя таблицы и словарь с яйчейками, которые нужно объединить(попарно)   
    """
    for i in range(1, len(cells), 2):
        table.cell(cells[str(i)][0], cells[str(i)][1]).merge(table.cell(cells[str(i+1)][0], cells[str(i+1)][1]))

        
def filling_table_heads_all(table_name, list_heads, font_sz=10):
    """
    Универсальная функция заполнения заголовков в таблицах
    Для заполнения заголовкав используется:
    - список для заголовков соответствующей таблицы в settings
    - имя таблицы
    - аргумент для задания размера шрифта font_sz
    """
    unique = list() # - список для формирования набора ячеек построково
    cells_list = list() # - список самих ячеек, понадобится для заполнения
    for row in table_name.rows:
        for cell in row.cells:
            tc = cell._tc # - доступ к свойствам ячеки
            cell_loc = (tc.top, tc.bottom, tc.left, tc.right) # - координаты ячеек по номеру границы
            if cell_loc not in unique:
                unique.append(cell_loc) # - формирования списка ячеек
                cells_list.append(cell)
    for i, item in enumerate(list_heads):
        p = cells_list[i].paragraphs[0]
        p.add_run(item)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.runs[0].font.size = Pt(font_sz)


def func_def_test_by_program(dict_tests: dict, union_tests_dict: dict) -> list:
    """
    Функция, которая читает испытание из программы и сопостовляет ему формулировку для протокола
    dict_tests - словарь с названиями испытаний из ПМИ
    union_tests_dict - словарь для сопоставления названия программы и формулировки в протоколе
    """
    list_form_protocols = []
    for value in dict_tests.values():
        if value in union_tests_dict.keys() and type(union_tests_dict[value]) == list:
            for i in union_tests_dict[value]:
                list_form_protocols.append(i)
        elif value in union_tests_dict.keys():
            list_form_protocols.append(union_tests_dict[value])
        else:
            pass
    return list_form_protocols

                         
def func_def_category_tests(list_form_protocols: list) -> dict:
    """Функция, которая определяет категорию испытаний и возвращает словарь, где
    ключ - номер категории, а значение список испытаний
    list_form_protocols - список с формулировками для протокола, выбранными из ПМИ
    list_records_test - словарь с категориями в виде числа и списками со всеми формулировками
    dict_for_temp_protocol - новый словарь сформированный из первых двух с ключем категорий и со значением
    в виде списка испытаний выбранных из программы
    """
    dict_for_temp_protocol = {}
    for test in list_form_protocols:
        for key, value in list_records_test.items():
            if test in value:
                dict_for_temp_protocol[key] = dict_for_temp_protocol.get(key, []) + [test]
            else:
                pass
    
    return dict_for_temp_protocol
        
    
def func_calculate_cells(row) -> int:
    """Функция считающая количество ячеек в строке"""
    counted_cells = set()
    for cell in row:
        tc = cell._tc
        cell_loc = (tc.top, tc.bottom, tc.left, tc.right)
        counted_cells.add(cell_loc)
    return(len(counted_cells))

def read_json_file(path) -> dict:
    """Читает json-файл и возвращает словарь"""
    with open(path, encoding='utf-8') as f:
        template = json.load(f)
        
    return template
        

svfile = read_json_file("D:\My_projects\LabReports\meta.json")    


# for chapter_num, chapter_text in svfile.items():
#     if chapter_num != "8":
#         for chapter_name, chapter_value in chapter_text.items():
#             print(chapter_num, chapter_name)
#             print(chapter_value)
#     else: 
#         break

# # for method, method_name in svfile["7"]['Методы испытаний'].items():
# #     print(method, str('"' + method_name + '";'))  




