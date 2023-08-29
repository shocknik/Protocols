from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.run import Run


def add_mergefield(field_name: str, **kwargs) -> Run:
    '''
    Add mergefield in docx.text.run

    Usage: add_mergefield(
        'str',
        before='text',
        after='text'
    )

    :param field_name: the name of new field
    :param kwargs:
        run = run where you need to set mergfiled;
        format = one of |Upper|, |Lower|, |FirstCap|, |TitleCase|;
        before = text before field;
        after = text after field;
        mapped = True - mapped field;
        vertical = True - vertical format
    :return: Run with added mergfiled
    '''
    if 'run' in kwargs:
        run = kwargs['run']
    else:
        run = Document().add_paragraph('').add_run()._r

    field_option = ''
    field = ''

    if field_name[0:1] == '«' and field_name[-1:] == '»':
        field_option = f' MERGEFIELD ' + field_name[1:-1]
        field = field_name
    else:
        field_option = f' MERGEFIELD ' + field_name
        field = '«' + field_name + '»'

    ordered_kwargs = {}
    if 'format' in kwargs: ordered_kwargs['format'] = kwargs['format']
    if 'before' in kwargs: ordered_kwargs['before'] = kwargs['before']
    if 'after' in kwargs: ordered_kwargs['after'] = kwargs['after']
    if 'mapped' in kwargs: ordered_kwargs['mapped'] = kwargs['mapped']
    if 'vertical' in kwargs: ordered_kwargs['vertical'] = kwargs['vertical']

    for key, value in ordered_kwargs.items():
        if key == 'format':
            if value == 'Upper':
                field_option += f' \* Upper'
            if value == 'Lower':
                field_option += f' \* Lower'
            if value == 'FirstCap':
                field_option += f' \* FirstCap'
            if value == 'TitleCase':
                field_option += f' \* Caps'
        if key == 'before':
            field_option += f' \\b ' + value
        if key == 'after':
            field_option += f' \\f ' + value
        if key == 'mapped' and value == True:
            field_option += f' \\m'
        if key == 'vertical' and value == True:
            field_option += f' \\v'

    field_option += f' \* MERGEFORMAT '

    ordered_kwargs = {}
    if 'before' in kwargs: ordered_kwargs['before'] = kwargs['before']
    if 'after' in kwargs: ordered_kwargs['after'] = kwargs['after']
    if 'format' in kwargs: ordered_kwargs['format'] = kwargs['format']

    for key, value in ordered_kwargs.items():
        if key == 'before':
            field = value + ' ' + field
        if key == 'after':
            field = field + ' ' + value
        if key == 'format':
            if value == 'Upper':
                field = field.upper()
            if value == 'Lower':
                field = field.lower()
            if value == 'FirstCap':
                field = field.capitalize()
            if value == 'TitleCase':
                old_field = field
                field = ''
                for str in old_field.split():
                    field += str.capitalize() + ' '
                field = field.strip()

    # <w:fldSimple w:instr=" MERGEFIELD $offerNumber \* Upper \b asd \* MERGEFORMAT ">
    # <w:r>
    # <w:t>ASD «$OFFERNUMBER»</w:t>
    # </w:r>
    # </w:fldSimple>
    fld = create_element('w:fldSimple', run)
    create_attribute(fld, 'w:instr', field_option)
    obj = create_element('w:r', fld)
    obj = create_element('w:t', obj)
    obj.text = field

    return run

def create_element(name:str, parent=None):
    '''
    Create new object in XML tree.

    :param name: type name of new object
    :param parent: obj created by OxmlElement()
    :return: created Object OR created child Object
    '''
    sub_obj = OxmlElement(name)
    if parent is not None:
        try:
            parent.append(sub_obj)
            return sub_obj
        except Exception:
            print('oops')
    else:
        return sub_obj


def create_attribute(element, name, value):
    element.set(qn(name), value)