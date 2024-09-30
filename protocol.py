import os
from setting import test_center, ratification, title_test_center, adress, cells_union, \
                    list_head_test_table
from borders import set_cell_border
from numpage import add_page_number
from docx import Document
from docx.oxml import OxmlElement, ns
from docx.shared import Inches, Cm, Mm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION_START, WD_ORIENTATION, WD_HEADER_FOOTER_INDEX
from datetime import datetime
from backend_read import get_cable_mark,\
    get_specifications,\
    get_name_specifications,\
    get_list_text_par,\
    create_marker_list,\
    change_font, border_form,\
    func_union_cells,\
    get_list_par_from_tables,\
    filling_table_heads_all,\
    get_list_requarements,\
    get_list_methods,\
    table_inner_border_vertical,\
    func_def_test_by_program, \
    read_json_file, \
    border_form
    
class Protocol:
    
    def __init__(self, path, path_json, test_center) -> None:
        self.path = path
        self.test_center = test_center
        self.path_json = path_json
        
        
    def create_new_file(self):
        """Метод создания нового файла"""
        doc = Document()
        """Задание альбомной формы, шрифта и отступов"""
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        title = doc.sections[0]
        title.orientation = WD_ORIENTATION.LANDSCAPE
        title.page_height = Cm(21.0)
        title.page_width = Cm(29.7)
        """формирование отступов полей"""
        title.left_margin = Mm(15)
        title.right_margin = Mm(13)
        title.top_margin = Mm(15)
        title.bottom_margin = Mm(13)
        doc.save(self.path)
        
    def create_title_list(self):
        """Метод создания титульного листа"""
        doc = Document(self.path)
        if self.test_center == 'KT':
            i = 0
        elif self.test_center =='SK':
            i = 1
        """Создание таблицы для рамки листа"""
        title_table = doc.add_table(rows=7, cols=3)
        cell_header = title_table.cell(0, 1)
        """Название Испытательного центра"""
        chp = cell_header.paragraphs[0]
        chp.add_run(title_test_center[i]).font.size = Pt(14)
        chp.runs[0].bold = True
        chp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell_header.width = Cm(23)
        chp.paragraph_format.space_after = Pt(0)
        set_cell_border(cell_header, bottom={"sz": 12, "val": "single", "color": "black", "space": "0"})
        if i == 0:
            cell_header_2 = title_table.cell(1, 1)
            chp2 = cell_header_2.paragraphs[0]
            chp2.add_run('ИСПЫТАТЕЛЬНЫЙ ЦЕНТР').bold = True
            chp2.paragraph_format.space_before = Pt(0)
            chp2.paragraph_format.space_after = Cm(1)
            chp2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif i == 1:
            cell_header_2 = title_table.cell(1, 0)
            cell_header_2.paragraphs[0].add_run().add_picture('logoSK.png', width=Inches(1.25))
        """Утверждающая надпись"""
        title_table.cell(1, 2).merge(title_table.cell(2, 2))
        cell_seo = title_table.cell(1, 2)
        cs = cell_seo.paragraphs[0]
        cs.add_run("УТВЕРЖДАЮ\n").bold = True
        cs.add_run(f'{ratification[i]}\n')
        cs.add_run(f'«___» _____________ {datetime.now().year}')
        cs.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        """Адрес осуществления деятельности"""
        title_table.cell(2, 0).merge(title_table.cell(2, 1))
        cell_adress = title_table.cell(2, 0)
        ca = cell_adress.paragraphs[0]
        ca.add_run(adress[i])
        ca.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        ca.paragraph_format.line_spacing = Pt(12)
        cell_adress.width = Cm(15)
        """Номер протокола"""
        cell_prot = title_table.cell(3, 1)
        cpp = cell_prot.paragraphs[0]
        cpp.add_run("ПРОТОКОЛ №            \n").bold = True
        cpp.add_run("от              \nприемочных испытаний")
        cpp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cpp.paragraph_format.space_after = Pt(0)
        """Описание протокола испытаний"""
        cell_discription = title_table.cell(4, 1)
        cell_discription.width = Cm(30)
        cd = cell_discription.paragraphs[0]
        cd.add_run("кабеля                                      марки ")
        cd.add_run("{},\n".format(get_cable_mark()[0])).bold = True
        cd.add_run("изготовленного ООО НПП «СПЕЦКАБЕЛЬ» на соответствие требованиям {} {}".\
            format(get_specifications(), get_name_specifications()))
        cd.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_HI
        cd.paragraph_format.line_spacing = Pt(12)
        cd.paragraph_format.space_after = Pt(6)
        """Доп информация о протоколе"""
        title_table.cell(5, 0).merge(title_table.cell(5, 1))
        title_table.cell(5, 1).merge(title_table.cell(5, 2))
        cell_remarks = title_table.cell(5, 0)
        crp = cell_remarks.paragraphs[0]
        crp.add_run(f'1 Листов всего: ')
        add_page_number(crp, 'NUMPAGES')
        crp.add_run(f'\n2 Результаты испытаний распространяются только на предоставленный (е) заказчиком образец (цы).\n\
3 Протокол испытаний не может быть частично или полностью воспроизведен без письменного разрешения Испытательного центра."')
        crp.paragraph_format.space_before = Cm(2)
        """Надпись Москва 2023"""
        cell_moscow = title_table.cell(6, 1)
        cym = cell_moscow.paragraphs[0]
        cym.add_run(f"Москва\n{datetime.now().year}")
        cym.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cym.paragraph_format.space_before = Cm(1)
        doc.sections[0].different_first_page_header_footer = True       
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        doc.save(self.path)
    
    def create_two_list(self):
        doc = Document(self.path)
        """Формирование колонтитулов"""
        doc.sections[1].different_first_page_header_footer = False
        header_2 = doc.sections[1].header
        header_2.paragraphs[0].text = 'Протокол №'
        header_2.add_paragraph().add_run('Лист ')
        add_page_number(header_2.paragraphs[1], 'PAGE')
        header_2.paragraphs[1].paragraph_format.space_after = Pt(0)
        header_2.add_paragraph().add_run('Всего листов ')
        add_page_number(header_2.paragraphs[2], 'NUMPAGES')
        header_2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        header_2.paragraphs[1].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        header_2.paragraphs[2].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        """Заполнение по данным из json"""
        data = read_json_file(self.path_json)
        for key, value in data.items():
            is_dict = True # для проверки, является ли элемент словарем, если да, то цикл распаковки продолжается, если нет, выводится содержимое строки
            while is_dict is True and int(key) < 8:
                if isinstance(value, dict):
                    for key_1, value_1 in value.items():
                        doc.add_paragraph().add_run(key).bold = True
                        doc.paragraphs[int(key)].add_run(" ")
                        doc.paragraphs[int(key)].add_run(key_1).bold = True
                        doc.paragraphs[int(key)].add_run("\n")
                        if isinstance(value_1, dict):
                            for key_2, value_2 in value_1.items():
                                doc.paragraphs[int(key)].add_run(key_2)
                                doc.paragraphs[int(key)].add_run(" ")
                                if int(key) == 7: doc.paragraphs[int(key)].add_run("«")
                                doc.paragraphs[int(key)].add_run(value_2)
                                if int(key) == 7: doc.paragraphs[int(key)].add_run("»;")
                                doc.paragraphs[int(key)].add_run("\n")
                            is_dict = False
                        else:
                            doc.paragraphs[int(key)].add_run(value_1)
                            doc.paragraphs[int(key)].add_run("\n")
                            is_dict = False
                else:
                    is_dict = False
        doc.save(self.path)

    def create_results_table(self):
        doc = Document(self.path)
        data = read_json_file(self.path_json)
        for key, value in data.items():
            if int(key) == 8:
                doc.add_paragraph().add_run(str(key + " " + list(value.keys())[0])).bold = True
            else:
                pass
        test_table = doc.add_table(rows=2, cols=7)
        test_table.cell(0, 0).width = Cm(6)
        table_inner_border_vertical(2, 7, test_table, sz=6, vert=False)
        border_form(2, 7, test_table, border="double", sz=6)
        func_union_cells(test_table, **cells_union)
        filling_table_heads_all(test_table, list_head_test_table)
        test_table.add_row().cells
        doc.save(self.path)
        
        
        





       
obj = Protocol(path = 'D:\\My_projects\\Protoсols\\tests_24.docx',\
                path_json="D:\My_projects\LabReports\meta.json",\
                test_center=test_center[1])
obj.create_new_file()
obj.create_title_list()
obj.create_two_list()
obj.create_results_table()

os.startfile("D:\\My_projects\\Protoсols\\tests_24.docx")



