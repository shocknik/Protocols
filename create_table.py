"""Файл содержащий функцию создания таблицы испытаний"""

from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.text import font
from docx.oxml.ns import qn
from borders import set_cell_border
from backend_read import border_form,\
    border_around_cell,\
        table_inner_border_vertical,\
            func_calculate_cells,\
                list_for_part_row

class Test_Table:
    """Класс таблицы"""
    
    def __init__(self, table_name, test_record, requrement, method, mean_req, limit) -> None:
        self.table_name = table_name
        self.test_record = test_record
        self.requrement = requrement
        self.method = method
        self.mean_req = mean_req
        self.limit = limit
        
        
    def row_for_navigation(self):
        """
        Метод, создающий строку номирации столбцов, которая переносится
        на каждый новый лист
        """
        cell = self.table_name.add_row().cells
        for i in range(1, func_calculate_cells(cell)+1):
            cell[i-1].text = str(i)
            cell[i-1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell[i-1].paragraphs[0].runs[0].font.size = Pt(10)
            cell[i-1].paragraphs[0].paragraph_format.space_after = Pt(0)
            set_cell_border(cell[i-1], end={"sz": 6, "val": 'single', "space": "0"})
        set_cell_border(cell[0], start={"sz": 6, "val": 'double', "space": "0"})
        set_cell_border(cell[func_calculate_cells(cell)-1], end={"sz": 6, "val": 'double', "space": "0"})

    
    def title_row(self, num, name):
        """
        Метод, который создает строку раздела испытаний
        num - порядковый номер раздела по ходу протокола
        name - наименование категории испытаний
        """
        cell = self.table_name.add_row().cells # Создание строки в таблице
        cell[0].merge(cell[6]) # объединений всех ячеек
        cell[0].text = str(num) + " " + str(name) # вставка текста
        cell[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # выравнивание текста по центру
        cell[0].paragraphs[0].runs[0].bold = True # выделение жирным
        cell[0].paragraphs[0].paragraph_format.space_after = Pt(0) # убрать отступ внизу
        cell[0].paragraphs[0].runs[0].font.size = Pt(12) # установить размер шрифта
        border_around_cell(cell[0])
        
    def title_row_from_datafrme(self, text):
        """Метод, который создает строку раздела испытаний по данным,
        полученным из датафрейма
        """
        cell = self.table_name.add_row().cells
        cell[0].merge(cell[6])
        cell[0].text = str(text)
        cell[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER # выравнивание текста по центру
        cell[0].paragraphs[0].runs[0].bold = True # выделение жирным
        cell[0].paragraphs[0].paragraph_format.space_after = Pt(0) # убрать отступ внизу
        cell[0].paragraphs[0].runs[0].font.size = Pt(12) # установить размер шрифта
        border_around_cell(cell[0], border="double", sz=6)
        
    
    def create_simple_row(self, num, tag=True):
        """Метод, который создает единичную строку под простые испытания"""      
        cell = self.table_name.add_row().cells
        cell[0].text = str(num) + " "
        cell[0].text += str(self.test_record)
        cell[0].paragraphs[0].runs[0].font.size = Pt(10)
        for i, item in enumerate((self.requrement, self.method, self.mean_req, self.limit), start=1):
            cell[i].text = str(item)
            cell[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell[i].paragraphs[0].runs[0].font.size = Pt(10)
        for i in range(1, func_calculate_cells(cell)+1):
            set_cell_border(cell[i-1], end={"sz": 6, "val": 'single', "space": "0"})
            if tag is True:
                set_cell_border(cell[i-1], bottom={"sz": 6, "val": 'single', "space": "0"}) 
        set_cell_border(cell[0], start={"sz": 6, "val": 'double', "space": "0"})
        set_cell_border(cell[func_calculate_cells(cell)-1], end={"sz": 6, "val": 'double', "space": "0"})

    
    def create_sample_par_row(self, num, **kwargs):
        """Метод, который создает доп строку под параметры образца"""
        cell = self.table_name.add_row().cells
        cell[0].paragraphs[0].add_run(str(num)).font.size = Pt(10)
        cell[0].paragraphs[0].add_run(' Параметры образцов и условия испытания:').font.size = Pt(10)
        count_of_paragraphs = 1
        for value in kwargs.values():
            cell[0].add_paragraph().add_run(value).font.size = Pt(10)
            count_of_paragraphs += 1
        for i in range(0, count_of_paragraphs):
            cell[0].paragraphs[i].paragraph_format.space_after = Pt(0)
            cell[0].paragraphs[i].paragraph_format.space_before = Pt(0)
        for i in range(1, func_calculate_cells(cell)+1):
            set_cell_border(cell[i-1], end={"sz": 6, "val": 'single', "space": "0"})
        set_cell_border(cell[0], start={"sz": 6, "val": 'double', "space": "0"})
        set_cell_border(cell[func_calculate_cells(cell)-1], end={"sz": 6, "val": 'double', "space": "0"})

    
    def create_validity_criteria(self, num, **kwargs):
        """Метод, который создает доп строку под критерии-годности"""
        cell = self.table_name.add_row().cells
        cell[0].paragraphs[0].add_run(str(num)).font.size = Pt(10)
        cell[0].paragraphs[0].add_run(' Критерии годности:').font.size = Pt(10)
        count_of_paragraphs = 1
        for value in kwargs.values():
            cell[0].add_paragraph().add_run(value[0]).font.size = Pt(10)
            cell[3].add_paragraph().add_run(value[1]).font.size = Pt(10)
            count_of_paragraphs += 1
        for i in range(0, count_of_paragraphs):
            cell[0].paragraphs[i].paragraph_format.space_after = Pt(0)
            cell[0].paragraphs[i].paragraph_format.space_before = Pt(0)
            cell[3].paragraphs[i].paragraph_format.space_after = Pt(0)
            cell[3].paragraphs[i].paragraph_format.space_before = Pt(0)
        for i in range(1, func_calculate_cells(cell)+1):
            set_cell_border(cell[i-1], end={"sz": 6, "val": 'single', "space": "0"})
            set_cell_border(cell[i-1], bottom={"sz": 6, "val": 'single', "space": "0"})
        set_cell_border(cell[0], start={"sz": 6, "val": 'double', "space": "0"})
        set_cell_border(cell[func_calculate_cells(cell)-1], end={"sz": 6, "val": 'double', "space": "0"})